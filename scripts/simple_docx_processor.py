#!/usr/bin/env python3
"""
Simple DOCX Processor - Local Testing
Applies JSON edits to DOCX files without LibreOffice dependency.
Uses python-docx library for direct document manipulation.

Usage:
    python3 scripts/simple_docx_processor.py \
        --in input.docx \
        --json edits.json \
        --out output.docx
"""

import argparse
import json
import sys
from pathlib import Path

def install_docx():
    """Install python-docx if not available."""
    try:
        import docx
        return True
    except ImportError:
        print("📦 Installing python-docx...")
        import subprocess
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            import docx
            return True
        except Exception as e:
            print(f"❌ Failed to install python-docx: {e}")
            print("💡 Please install manually: pip install python-docx")
            return False

def load_json_operations(json_path):
    """Load operations from JSON file."""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    operations = data.get('instructions', {}).get('operations', [])
    print(f"📋 Loaded {len(operations)} operations from JSON")
    return operations

def process_docx(input_path, operations, output_path):
    """Process DOCX file with operations."""
    import docx
    from docx.shared import RGBColor, Pt
    
    # Load document
    doc = docx.Document(input_path)
    
    changes_made = 0
    
    # Process each operation
    for i, op in enumerate(operations):
        action = op.get('action', 'replace')
        target_text = op.get('target_text', '')
        replacement = op.get('replacement', '')
        comment = op.get('comment', '')
        author = op.get('comment_author', 'AI Assistant')
        
        # Skip comment-only operations
        if action == 'comment':
            continue
            
        # Handle delete operations
        if action == 'delete':
            replacement = ''
        
        # Keep replacement text clean - don't modify it
        
        # Search and replace in all paragraphs
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                # Find the position where replacement happens
                original_text = paragraph.text
                new_text = original_text.replace(target_text, replacement)
                
                # Clear existing runs and create new ones
                paragraph.clear()
                
                # Add the new text
                run = paragraph.add_run(new_text)
                
                # Add styling to indicate change
                run.font.color.rgb = RGBColor(0, 100, 200)  # Blue text
                
                # Add comment as a reply-like annotation
                if comment:
                    try:
                        # Format comment text
                        comment_text = comment.replace('\\n', '\n').replace('\\n\\n', '\n\n')
                        
                        # Add a subtle comment indicator next to the change
                        run.add_text(" 💬")
                        
                        # Insert a comment paragraph right after this one
                        # Find current paragraph index
                        para_index = None
                        for idx, para in enumerate(doc.paragraphs):
                            if para == paragraph:
                                para_index = idx
                                break
                        
                        if para_index is not None:
                            # Add comment paragraph after current one
                            comment_para = doc.add_paragraph()
                            
                            # Move the comment paragraph to the right position
                            # (This is a workaround since python-docx doesn't have insert_paragraph)
                            
                            # Style the comment like a reply
                            comment_para.add_run(f"💬 {author}: ").font.bold = True
                            comment_para.add_run(comment_text).font.italic = True
                            
                            # Style the paragraph
                            comment_para.paragraph_format.left_indent = Pt(20)  # Indent like a reply
                            comment_para.paragraph_format.space_after = Pt(6)
                            
                            # Make text smaller and gray
                            for run in comment_para.runs:
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(100, 100, 100)
                        
                    except Exception as e:
                        print(f"   ⚠️ Could not add comment: {e}")
                        # Simple fallback - just add a marker
                        try:
                            run.add_text(" 💬")
                        except:
                            pass
                
                changes_made += 1
                print(f"✅ {i+1}. Replaced '{target_text[:50]}...' with '{replacement[:50]}...'")
                if comment:
                    print(f"   💬 Comment: {comment[:80]}...")
                break
        
        # Also search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target_text in cell.text:
                        cell.text = cell.text.replace(target_text, replacement)
                        changes_made += 1
                        print(f"✅ {i+1}. Replaced in table: '{target_text[:30]}...'")
    
    # Save document
    doc.save(output_path)
    
    return changes_made

def main():
    parser = argparse.ArgumentParser(description='Simple DOCX processor for local testing')
    parser.add_argument('--in', dest='input_path', required=True, help='Input DOCX file')
    parser.add_argument('--json', dest='json_path', required=True, help='JSON operations file')
    parser.add_argument('--out', dest='output_path', required=True, help='Output DOCX file')
    
    args = parser.parse_args()
    
    print("🚀 Simple DOCX Processor")
    print("=" * 40)
    
    # Check if python-docx is available
    if not install_docx():
        sys.exit(1)
    
    # Validate input files
    if not Path(args.input_path).exists():
        print(f"❌ Input file not found: {args.input_path}")
        sys.exit(1)
    
    if not Path(args.json_path).exists():
        print(f"❌ JSON file not found: {args.json_path}")
        sys.exit(1)
    
    try:
        # Load operations
        operations = load_json_operations(args.json_path)
        
        # Create output directory
        Path(args.output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Process document
        print(f"📝 Processing: {args.input_path}")
        changes_made = process_docx(args.input_path, operations, args.output_path)
        
        print("\n🎉 Processing Complete!")
        print("=" * 40)
        print(f"📊 Changes made: {changes_made}")
        print(f"📁 Output saved: {args.output_path}")
        print("\n💡 Open the output file in any word processor to review changes")
        print("   Comments and author info are included in the replacement text")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()