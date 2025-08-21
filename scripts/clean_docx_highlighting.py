#!/usr/bin/env python3
"""
DOCX Highlighting Cleaner - Remove all highlighted text from DOCX files

This utility removes all highlighting from DOCX files, which can interfere 
with automated processing systems.

Usage:
    python3 clean_docx_highlighting.py input.docx [output.docx]
    
    If output.docx is not specified, the input file will be overwritten.

Examples:
    python3 clean_docx_highlighting.py policy.docx policy_clean.docx
    python3 clean_docx_highlighting.py policy.docx  # overwrites policy.docx
"""

import os
import sys
import argparse
from pathlib import Path

def clean_docx_highlighting(input_path, output_path=None):
    """Remove all highlighting from a DOCX file."""
    try:
        import docx
        from docx.shared import RGBColor
        
        # If no output path specified, overwrite the input file
        if output_path is None:
            output_path = input_path
            
        doc = docx.Document(input_path)
        highlighting_removed = 0
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Remove highlight color
                try:
                    if run.font.highlight_color is not None:
                        run.font.highlight_color = None
                        highlighting_removed += 1
                except:
                    pass
                
                # Remove background color
                try:
                    if hasattr(run.font, 'fill'):
                        run.font.fill.fore_color.rgb = None
                except:
                    pass
        
        # Process tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                if run.font.highlight_color is not None:
                                    run.font.highlight_color = None
                                    highlighting_removed += 1
                            except:
                                pass
        
        # Save the cleaned document
        doc.save(output_path)
        
        return True, highlighting_removed
        
    except ImportError:
        return False, "python-docx package is required. Install with: pip install python-docx"
    except Exception as e:
        return False, f"Error cleaning DOCX highlighting: {e}"

def main():
    parser = argparse.ArgumentParser(description='Remove highlighting from DOCX files')
    parser.add_argument('input', help='Input DOCX file path')
    parser.add_argument('output', nargs='?', help='Output DOCX file path (optional, defaults to overwriting input)')
    parser.add_argument('--backup', action='store_true', help='Create a backup of the original file before cleaning')
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    
    # Validate input file
    if not input_path.exists():
        print(f"❌ Error: Input file not found: {input_path}")
        sys.exit(1)
    
    if not input_path.suffix.lower() == '.docx':
        print(f"❌ Error: Input file must be a .docx file: {input_path}")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path
    
    # Create backup if requested
    if args.backup and output_path == input_path:
        backup_path = input_path.with_suffix('.backup.docx')
        import shutil
        shutil.copy2(input_path, backup_path)
        print(f"📁 Backup created: {backup_path}")
    
    print(f"🎨 Cleaning highlighting from: {input_path}")
    if output_path != input_path:
        print(f"💾 Output will be saved to: {output_path}")
    else:
        print(f"💾 File will be overwritten in place")
    
    # Clean the highlighting
    success, result = clean_docx_highlighting(str(input_path), str(output_path))
    
    if success:
        if result > 0:
            print(f"✅ Successfully removed highlighting from {result} text runs")
        else:
            print(f"✅ No highlighting found - document was already clean")
        print(f"📄 Cleaned document saved: {output_path}")
    else:
        print(f"❌ Error: {result}")
        sys.exit(1)

if __name__ == "__main__":
    main()
