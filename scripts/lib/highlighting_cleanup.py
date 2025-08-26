"""
DOCX Document Utilities

This module provides utilities for working with DOCX documents:
- Removing highlighting from documents
- Extracting clean text content
- Processing headers, footers, and tables
"""

import warnings
from typing import Tuple, Optional

# Suppress docx warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)


def clean_docx_highlighting(input_path: str, output_path: Optional[str] = None) -> Tuple[bool, str]:
    """
    Remove all highlighting from a DOCX file.
    
    Args:
        input_path: Path to input DOCX file
        output_path: Path for output (defaults to overwriting input)
        
    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        import docx
        from docx.shared import RGBColor
        
        # If no output path specified, overwrite the input file
        if output_path is None:
            output_path = input_path
            
        doc = docx.Document(input_path)
        highlighting_removed = 0
        
        # Process all paragraphs with comprehensive highlighting removal
        for paragraph in doc.paragraphs:
            highlighting_removed += _clean_paragraph_highlighting(paragraph)
        
        # Process tables with comprehensive highlighting removal
        for table in doc.tables:
            highlighting_removed += _clean_table_highlighting(table)
        
        # Process headers and footers
        for section in doc.sections:
            highlighting_removed += _clean_section_highlighting(section)
        
        # Save the cleaned document
        doc.save(output_path)
        
        if highlighting_removed > 0:
            print(f"🎨 Removed highlighting from {highlighting_removed} text runs in DOCX")
        
        return True, f"Cleaned {highlighting_removed} highlighted sections"
        
    except ImportError:
        return False, "python-docx not available for highlighting removal"
    except Exception as e:
        return False, f"Error cleaning DOCX highlighting: {e}"


def extract_docx_content(file_path: str, filter_highlighted: bool = True) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        filter_highlighted: Whether to exclude highlighted text
        
    Returns:
        Extracted text content
    """
    try:
        import docx
        doc = docx.Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            # Extract text while optionally filtering out highlighted content
            clean_text = ""
            for run in paragraph.runs:
                # Check if the run has highlighting
                is_highlighted = False
                if filter_highlighted:
                    is_highlighted = _is_run_highlighted(run)
                
                # Only add text if it's not highlighted (or we're not filtering)
                if not is_highlighted:
                    clean_text += run.text
            
            # Only add non-empty paragraphs
            if clean_text.strip():
                content.append(clean_text)
        
        filtered_content = '\n'.join(content)
        status = "highlighted text removed" if filter_highlighted else "all text included"
        print(f"📄 DOCX loaded: {len(filtered_content)} characters ({status})")
        return filtered_content
        
    except ImportError:
        return f"[DOCX FILE: {file_path} - Install python-docx to read content]"
    except Exception as e:
        raise Exception(f"Error reading DOCX file {file_path}: {e}")


def _clean_paragraph_highlighting(paragraph) -> int:
    """Clean highlighting from a single paragraph."""
    highlighting_removed = 0
    
    # Clear paragraph-level highlighting first
    try:
        if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
            # Remove paragraph shading elements
            para_shading = paragraph._element.pPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            for shd in para_shading:
                paragraph._element.pPr.remove(shd)
                highlighting_removed += 1
    except:
        pass
    
    for run in paragraph.runs:
        highlighting_removed += _clean_run_highlighting(run)
    
    return highlighting_removed


def _clean_table_highlighting(table) -> int:
    """Clean highlighting from a table."""
    highlighting_removed = 0
    
    for row in table.rows:
        for cell in row.cells:
            # Clear cell-level shading
            try:
                if hasattr(cell._element, 'tcPr') and cell._element.tcPr is not None:
                    cell_shading = cell._element.tcPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    for shd in cell_shading:
                        cell._element.tcPr.remove(shd)
                        highlighting_removed += 1
            except:
                pass
            
            for paragraph in cell.paragraphs:
                highlighting_removed += _clean_paragraph_highlighting(paragraph)
    
    return highlighting_removed


def _clean_section_highlighting(section) -> int:
    """Clean highlighting from headers and footers in a section."""
    highlighting_removed = 0
    
    # Process all types of headers
    for header_type in ['first_page_header', 'even_page_header', 'header']:
        try:
            header = getattr(section, header_type)
            if header is not None:
                for paragraph in header.paragraphs:
                    highlighting_removed += _clean_paragraph_highlighting(paragraph)
                
                # Process tables in headers if any
                for table in header.tables:
                    highlighting_removed += _clean_table_highlighting(table)
                
                # Process graphics/images in headers
                highlighting_removed += _clean_graphics_highlighting(header)
        except:
            pass
    
    # Process all types of footers
    for footer_type in ['first_page_footer', 'even_page_footer', 'footer']:
        try:
            footer = getattr(section, footer_type)
            if footer is not None:
                for paragraph in footer.paragraphs:
                    highlighting_removed += _clean_paragraph_highlighting(paragraph)
                
                # Process graphics/images in footers
                highlighting_removed += _clean_graphics_highlighting(footer)
        except:
            pass
    
    return highlighting_removed


def _clean_run_highlighting(run) -> int:
    """Clean highlighting from a single text run."""
    highlighting_removed = 0
    
    # Method 1: Remove highlight color (main highlighting property)
    try:
        if run.font.highlight_color is not None:
            run.font.highlight_color = None
            highlighting_removed += 1
    except:
        pass
    
    # Method 2: Remove font fill/background colors
    try:
        if hasattr(run.font, 'fill') and run.font.fill is not None:
            if hasattr(run.font.fill, 'fore_color') and run.font.fill.fore_color is not None:
                run.font.fill.fore_color.rgb = None
                highlighting_removed += 1
            # Clear any fill type
            run.font.fill.solid()
    except:
        pass
    
    # Method 3: Clear character shading/background (XML level)
    try:
        if hasattr(run._element, 'rPr') and run._element.rPr is not None:
            # Remove ALL shading elements
            shading_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            for shd in shading_elements:
                run._element.rPr.remove(shd)
                highlighting_removed += 1
    except:
        pass
    
    # Method 4: Remove character spacing and position (sometimes contains highlighting info)
    try:
        if hasattr(run._element, 'rPr') and run._element.rPr is not None:
            # Remove highlight-related XML properties
            highlight_props = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
            for prop in highlight_props:
                run._element.rPr.remove(prop)
                highlighting_removed += 1
    except:
        pass
    
    # Method 5: Clear any w:color background attributes
    try:
        if hasattr(run._element, 'rPr') and run._element.rPr is not None:
            color_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            for color_elem in color_elements:
                # Remove background color attributes
                if color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                    del color_elem.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill']
                    highlighting_removed += 1
    except:
        pass
    
    return highlighting_removed


def _clean_graphics_highlighting(header_or_footer) -> int:
    """Clean highlighting from graphics/images in headers or footers."""
    highlighting_removed = 0
    
    try:
        # Access the underlying XML element to find inline shapes (graphics)
        header_element = header_or_footer._element
        
        # Find all drawing elements (graphics/images) in the header/footer
        # Look for w:drawing elements which contain images
        drawings = header_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        
        for drawing in drawings:
            # Look for any highlighting/shading properties on the drawing's parent run
            parent_run = drawing.getparent()
            if parent_run is not None:
                # Find run properties
                run_props = parent_run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if run_props is not None:
                    # Remove highlighting elements
                    highlight_elements = run_props.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                    for highlight in highlight_elements:
                        run_props.remove(highlight)
                        highlighting_removed += 1
                    
                    # Remove shading elements
                    shading_elements = run_props.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    for shading in shading_elements:
                        run_props.remove(shading)
                        highlighting_removed += 1
        
        if highlighting_removed > 0:
            print(f"🎨 Removed highlighting from {highlighting_removed} graphics in header/footer")
            
    except Exception as e:
        # Not critical - continue without error
        pass
    
    return highlighting_removed


def _is_run_highlighted(run) -> bool:
    """Check if a text run has highlighting applied."""
    try:
        # Check for highlight color (python-docx way)
        if run.font.highlight_color is not None:
            return True
    except:
        pass
    
    # Also check for background color highlighting
    try:
        if hasattr(run.font, 'fill') and run.font.fill.fore_color.rgb is not None:
            return True
    except:
        pass
    
    return False
