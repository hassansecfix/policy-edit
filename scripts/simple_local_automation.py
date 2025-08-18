#!/usr/bin/env python3
"""
Simple Local Automation - Complete policy customization without LibreOffice UNO

This script provides a simple alternative to the full automation that works entirely locally:
1. Converts Excel questionnaire to CSV (if needed)
2. Calls Claude Sonnet 4 to generate edits JSON
3. Applies edits directly to DOCX using python-docx (no LibreOffice required)

Usage:
    python3 scripts/simple_local_automation.py \
        --policy data/policy.docx \
        --questionnaire data/questionnaire.csv \
        --output-name "customized_policy" \
        --api-key YOUR_CLAUDE_API_KEY

Requirements:
    pip install python-docx requests anthropic

Environment Variables:
    CLAUDE_API_KEY: Your Anthropic Claude API key
"""

import os
import sys
import argparse
import json
import tempfile
import requests
from pathlib import Path

# Import our AI processor
sys.path.append(str(Path(__file__).parent))

def apply_edits_with_python_docx(docx_path, edits_json, output_path):
    """Apply edits to DOCX using python-docx library (simple replacements only)."""
    try:
        import docx
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        return False

    try:
        # Load the document
        doc = docx.Document(docx_path)
        
        # Load the edits
        with open(edits_json, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        operations = data.get('instructions', {}).get('operations', [])
        metadata = data.get('metadata', {})
        
        print(f"📝 Applying {len(operations)} operations to document...")
        
        # Track changes made
        changes_made = 0
        
        for op in operations:
            action = op.get('action', 'replace')
            target_text = op.get('target_text', '')
            replacement = op.get('replacement', '')
            comment = op.get('comment', '')
            author = op.get('comment_author', 'AI Assistant')
            
            if action == 'replace_with_logo':
                # Handle logo replacement
                logo_path = metadata.get('logo_path')
                if not logo_path:
                    # Try to download from questionnaire
                    logo_path = download_logo_from_questionnaire()
                
                if not logo_path and os.path.exists('data/company_logo.png'):
                    logo_path = 'data/company_logo.png'
                
                if logo_path and os.path.exists(logo_path):
                    # Replace in headers first
                    for section in doc.sections:
                        header = section.header
                        for paragraph in header.paragraphs:
                            if target_text in paragraph.text:
                                # Clear the paragraph and add the logo
                                paragraph.clear()
                                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                                try:
                                    # Get logo dimensions
                                    width = metadata.get('logo_width_mm', 35)
                                    height = metadata.get('logo_height_mm', 0)
                                    
                                    if width:
                                        width_inches = width / 25.4  # Convert mm to inches
                                        if height and height > 0:
                                            height_inches = height / 25.4
                                            run.add_picture(logo_path, width=Inches(width_inches), height=Inches(height_inches))
                                        else:
                                            run.add_picture(logo_path, width=Inches(width_inches))
                                    else:
                                        run.add_picture(logo_path)
                                    
                                    print(f"🖼️  Replaced '{target_text}' with logo in header")
                                    changes_made += 1
                                except Exception as e:
                                    print(f"⚠️  Could not insert logo: {e}")
                                    # Fallback: replace with text
                                    paragraph.text = f"[LOGO: {os.path.basename(logo_path)}]"
                                    changes_made += 1
                
            elif action == 'replace':
                # Handle text replacement
                replaced = False
                
                # Replace in main document
                for paragraph in doc.paragraphs:
                    if target_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(target_text, replacement)
                        print(f"✅ Replaced '{target_text[:50]}...' with '{replacement[:50]}...'")
                        changes_made += 1
                        replaced = True
                
                # Replace in headers/footers
                for section in doc.sections:
                    for header_paragraph in section.header.paragraphs:
                        if target_text in header_paragraph.text:
                            header_paragraph.text = header_paragraph.text.replace(target_text, replacement)
                            changes_made += 1
                            replaced = True
                    
                    for footer_paragraph in section.footer.paragraphs:
                        if target_text in footer_paragraph.text:
                            footer_paragraph.text = footer_paragraph.text.replace(target_text, replacement)
                            changes_made += 1
                            replaced = True
                
                # Replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if target_text in cell.text:
                                cell.text = cell.text.replace(target_text, replacement)
                                changes_made += 1
                                replaced = True
                
                if not replaced:
                    print(f"⚠️  Could not find text to replace: '{target_text[:50]}...'")
                    
            elif action == 'delete':
                # Handle text deletion
                deleted = False
                
                # Delete from main document
                for paragraph in doc.paragraphs:
                    if target_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(target_text, '')
                        print(f"🗑️  Deleted '{target_text[:50]}...'")
                        changes_made += 1
                        deleted = True
                
                if not deleted:
                    print(f"⚠️  Could not find text to delete: '{target_text[:50]}...'")
            
            elif action == 'comment':
                print(f"💬 Comment operation (not applied in simple mode): {comment[:50]}...")
        
        # Save the document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        print(f"\n🎉 SUCCESS! Applied {changes_made} changes")
        print(f"📁 Saved to: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error applying edits: {e}")
        return False

def download_logo_from_questionnaire():
    """Try to download logo from questionnaire URL."""
    try:
        questionnaire_path = 'data/secfix_questionnaire_responses_consulting.csv'
        if os.path.exists(questionnaire_path):
            with open(questionnaire_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if 'https://' in content and 'logo' in content.lower():
                    import re
                    urls = re.findall(r'https://[^\s,;]+', content)
                    for url in urls:
                        if any(ext in url.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif', 'image']):
                            # Download the logo
                            response = requests.get(url, stream=True, timeout=10)
                            if response.status_code == 200:
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                                for chunk in response.iter_content(chunk_size=8192):
                                    temp_file.write(chunk)
                                temp_file.close()
                                print(f"🌐 Downloaded logo from URL: {url}")
                                return temp_file.name
    except Exception as e:
        print(f"⚠️  Could not download logo: {e}")
    return None

def generate_edits_with_ai(policy_path, questionnaire_csv, prompt_path, policy_instructions_path, output_json, api_key):
    """Generate JSON instructions using AI."""
    import subprocess
    
    print("🧠 Generating JSON instructions with Claude Sonnet 4...")
    cmd = [
        'python3', 'scripts/ai_policy_processor.py',
        '--policy', policy_path,
        '--questionnaire', questionnaire_csv,
        '--prompt', prompt_path,
        '--policy-instructions', policy_instructions_path,
        '--output', output_json,
        '--api-key', api_key
    ]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        error_msg = result.stderr.strip() if result.stderr.strip() else result.stdout.strip()
        print(f"❌ AI generation failed: {error_msg}")
        return False
    
    print("✅ AI generation completed")
    return True

def convert_xlsx_to_csv(xlsx_path, csv_path):
    """Convert Excel questionnaire to CSV."""
    import subprocess
    
    print("📊 Converting Excel to CSV...")
    cmd = ['python3', 'scripts/xlsx_to_csv_converter.py', xlsx_path, csv_path]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        error_msg = result.stderr.strip() if result.stderr.strip() else result.stdout.strip()
        print(f"❌ Excel conversion failed: {error_msg}")
        return False
    
    print("✅ Excel conversion completed")
    return True

def main():
    parser = argparse.ArgumentParser(description='Simple Local Automation - Complete policy customization without LibreOffice')
    parser.add_argument('--policy', required=True, help='Path to policy DOCX file')
    parser.add_argument('--questionnaire', required=True, help='Path to questionnaire Excel/CSV file')
    parser.add_argument('--output-name', required=True, help='Base name for output files (e.g., "acme_policy")')
    parser.add_argument('--api-key', help='Claude API key (or set CLAUDE_API_KEY env var)')
    parser.add_argument('--logo', help='Optional path to company logo image (png/jpg)')
    parser.add_argument('--logo-width-mm', type=int, default=35, help='Logo width in millimeters (default: 35)')
    parser.add_argument('--logo-height-mm', type=int, default=0, help='Logo height in millimeters (0 = preserve aspect ratio)')
    
    args = parser.parse_args()
    
    # Get API key
    api_key = args.api_key or os.environ.get('CLAUDE_API_KEY')
    if not api_key:
        print("❌ Error: Claude API key required!")
        print("   Set CLAUDE_API_KEY environment variable or use --api-key")
        print("   Get your key from: https://console.anthropic.com/")
        sys.exit(1)
    
    print("🚀 Simple Local Automation Starting...")
    print("=" * 50)
    print(f"📋 Policy Document: {args.policy}")
    print(f"📊 Questionnaire: {args.questionnaire}")
    print(f"📝 Output Name: {args.output_name}")
    print(f"🤖 AI: Claude Sonnet 4")
    print(f"🔧 Mode: Local (python-docx)")
    print("=" * 50)
    
    # Create intermediate file paths
    questionnaire_csv = f"data/{args.output_name}_questionnaire.csv"
    edits_json = f"edits/{args.output_name}_edits.json"
    output_docx = f"build/{args.output_name}.docx"
    prompt_path = "data/prompt.md"
    policy_instructions_path = "data/updated_policy_instructions_v4.0.md"
    
    try:
        # Step 1: Convert Excel to CSV (if needed)
        if args.questionnaire.endswith(('.xlsx', '.xls')):
            print("\n📊 STEP 1: Converting Excel to CSV")
            if not convert_xlsx_to_csv(args.questionnaire, questionnaire_csv):
                sys.exit(1)
        else:
            # Already CSV, just use it
            questionnaire_csv = args.questionnaire
            print(f"\n📊 STEP 1: Using existing CSV: {questionnaire_csv}")
        
        # Step 2: Generate edits with AI
        print("\n🧠 STEP 2: Generating Edits with Claude Sonnet 4")
        if not generate_edits_with_ai(
            args.policy, questionnaire_csv, prompt_path, policy_instructions_path, 
            edits_json, api_key
        ):
            sys.exit(1)

        # Step 2.5: Inject logo metadata if provided
        if args.logo or args.logo_width_mm or args.logo_height_mm:
            try:
                with open(edits_json, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                data.setdefault('metadata', {})
                if args.logo:
                    data['metadata']['logo_path'] = args.logo
                if args.logo_width_mm:
                    data['metadata']['logo_width_mm'] = args.logo_width_mm
                if args.logo_height_mm:
                    data['metadata']['logo_height_mm'] = args.logo_height_mm
                with open(edits_json, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                print("🖼️  Injected logo metadata into edits JSON")
            except Exception as e:
                print(f"⚠️  Could not inject logo metadata: {e}")
        
        # Step 3: Apply edits locally
        print("\n📝 STEP 3: Applying Edits to Document")
        if not apply_edits_with_python_docx(args.policy, edits_json, output_docx):
            sys.exit(1)
        
        # Success summary
        print("\n🎉 AUTOMATION COMPLETE!")
        print("=" * 50)
        print("✅ Generated Files:")
        print(f"   📊 Questionnaire CSV: {questionnaire_csv}")
        print(f"   📋 JSON Instructions: {edits_json}")
        print(f"   📄 Final DOCX: {output_docx}")
        
        print("\n🔍 Next Steps:")
        print("1. Open the DOCX file in LibreOffice Writer or Microsoft Word")
        print("2. Review the applied changes")
        print("3. The document is ready for use!")
        
        print(f"\n📁 Output file: {output_docx}")
        print("🏆 Your policy has been customized!")
        
    except KeyboardInterrupt:
        print("\n⏹️  Automation cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
