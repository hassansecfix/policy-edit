#!/usr/bin/env python3
"""
AI Policy Processor - Automated CSV Generation with Claude Sonnet 4

This script takes a policy DOCX file, questionnaire CSV, and AI prompt,
then uses Claude Sonnet 4 to generate the perfect edits CSV file.

Usage:
    python3 ai_policy_processor.py \
        --policy data/policy.docx \
        --questionnaire data/questionnaire.csv \
        --prompt data/updated_policy_instructions_v4.0.md \
        --output edits/ai_generated_edits.csv \
        --api-key YOUR_CLAUDE_API_KEY

Environment Variables:
    CLAUDE_API_KEY: Your Anthropic Claude API key
"""

import os
import sys
import argparse
import csv
import re
import warnings
from pathlib import Path
import json

# Import anthropic only when needed (not when skipping API)
anthropic = None

# Suppress deprecation warnings for the Claude API
warnings.filterwarnings("ignore", category=DeprecationWarning)

def clean_docx_highlighting(input_path, output_path=None):
    """Remove all highlighting from a DOCX file."""
    try:
        import docx
        from docx.shared import RGBColor
        
        # If no output path specified, overwrite the input file
        if output_path is None:
            output_path = input_path
            
        doc = docx.Document(input_path)
        highlighting_removed = 0
        
        # Process all paragraphs with comprehensive highlighting removal
        for paragraph in doc.paragraphs:
            # Clear paragraph-level highlighting first
            try:
                if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                    # Remove paragraph shading elements
                    para_shading = paragraph._element.pPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    for shd in para_shading:
                        paragraph._element.pPr.remove(shd)
                        highlighting_removed += 1
            except:
                pass
            
            for run in paragraph.runs:
                # Method 1: Remove highlight color (main highlighting property)
                try:
                    if run.font.highlight_color is not None:
                        run.font.highlight_color = None
                        highlighting_removed += 1
                except:
                    pass
                
                # Method 2: Remove font fill/background colors
                try:
                    if hasattr(run.font, 'fill') and run.font.fill is not None:
                        if hasattr(run.font.fill, 'fore_color') and run.font.fill.fore_color is not None:
                            run.font.fill.fore_color.rgb = None
                            highlighting_removed += 1
                        # Clear any fill type
                        run.font.fill.solid()
                except:
                    pass
                
                # Method 3: Clear character shading/background (XML level)
                try:
                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                        # Remove ALL shading elements
                        shading_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                        for shd in shading_elements:
                            run._element.rPr.remove(shd)
                            highlighting_removed += 1
                except:
                    pass
                
                # Method 4: Remove character spacing and position (sometimes contains highlighting info)
                try:
                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                        # Remove highlight-related XML properties
                        highlight_props = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                        for prop in highlight_props:
                            run._element.rPr.remove(prop)
                            highlighting_removed += 1
                except:
                    pass
                
                # Method 5: Clear any w:color background attributes
                try:
                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                        color_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                        for color_elem in color_elements:
                            # Remove background color attributes
                            if color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                                del color_elem.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill']
                                highlighting_removed += 1
                except:
                    pass
        
        # Process tables with comprehensive highlighting removal
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Clear cell-level shading
                    try:
                        if hasattr(cell._element, 'tcPr') and cell._element.tcPr is not None:
                            cell_shading = cell._element.tcPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                            for shd in cell_shading:
                                cell._element.tcPr.remove(shd)
                                highlighting_removed += 1
                    except:
                        pass
                    
                    for paragraph in cell.paragraphs:
                        # Clear paragraph-level highlighting in tables
                        try:
                            if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                                para_shading = paragraph._element.pPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                for shd in para_shading:
                                    paragraph._element.pPr.remove(shd)
                                    highlighting_removed += 1
                        except:
                            pass
                        
                        for run in paragraph.runs:
                            # Apply all 5 methods to table text as well
                            try:
                                if run.font.highlight_color is not None:
                                    run.font.highlight_color = None
                                    highlighting_removed += 1
                            except:
                                pass
                            
                            try:
                                if hasattr(run.font, 'fill') and run.font.fill is not None:
                                    if hasattr(run.font.fill, 'fore_color') and run.font.fill.fore_color is not None:
                                        run.font.fill.fore_color.rgb = None
                                        highlighting_removed += 1
                                    run.font.fill.solid()
                            except:
                                pass
                            
                            try:
                                if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                                    # Remove shading elements
                                    shading_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                    for elem in shading_elements:
                                        run._element.rPr.remove(elem)
                                        highlighting_removed += 1
                                    
                                    # Remove highlight elements
                                    highlight_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                                    for elem in highlight_elements:
                                        run._element.rPr.remove(elem)
                                        highlighting_removed += 1
                            except:
                                pass
        
        # Process headers and footers (this was missing!)
        for section in doc.sections:
            # Process all types of headers
            for header_type in ['first_page_header', 'even_page_header', 'header']:
                try:
                    header = getattr(section, header_type)
                    if header is not None:
                        for paragraph in header.paragraphs:
                            # Clear paragraph-level highlighting in headers
                            try:
                                if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                                    para_shading = paragraph._element.pPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                    for shd in para_shading:
                                        paragraph._element.pPr.remove(shd)
                                        highlighting_removed += 1
                            except:
                                pass
                            
                            for run in paragraph.runs:
                                # Apply all 5 highlighting removal methods to header text
                                try:
                                    if run.font.highlight_color is not None:
                                        run.font.highlight_color = None
                                        highlighting_removed += 1
                                except:
                                    pass
                                
                                try:
                                    if hasattr(run.font, 'fill') and run.font.fill is not None:
                                        if hasattr(run.font.fill, 'fore_color') and run.font.fill.fore_color is not None:
                                            run.font.fill.fore_color.rgb = None
                                            highlighting_removed += 1
                                        run.font.fill.solid()
                                except:
                                    pass
                                
                                try:
                                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                                        # Remove shading elements
                                        shading_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                        for elem in shading_elements:
                                            run._element.rPr.remove(elem)
                                            highlighting_removed += 1
                                        
                                        # Remove highlight elements
                                        highlight_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                                        for elem in highlight_elements:
                                            run._element.rPr.remove(elem)
                                            highlighting_removed += 1
                                        
                                        # Remove color background attributes
                                        color_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                                        for color_elem in color_elements:
                                            if color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                                                del color_elem.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill']
                                                highlighting_removed += 1
                                except:
                                    pass
                        
                        # Process tables in headers if any
                        for table in header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    # Clear cell-level shading in header tables
                                    try:
                                        if hasattr(cell._element, 'tcPr') and cell._element.tcPr is not None:
                                            cell_shading = cell._element.tcPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                            for shd in cell_shading:
                                                cell._element.tcPr.remove(shd)
                                                highlighting_removed += 1
                                    except:
                                        pass
                                    
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            # Apply highlighting removal to table text in headers
                                            try:
                                                if run.font.highlight_color is not None:
                                                    run.font.highlight_color = None
                                                    highlighting_removed += 1
                                            except:
                                                pass
                except:
                    pass
            
            # Process all types of footers
            for footer_type in ['first_page_footer', 'even_page_footer', 'footer']:
                try:
                    footer = getattr(section, footer_type)
                    if footer is not None:
                        for paragraph in footer.paragraphs:
                            # Clear paragraph-level highlighting in footers
                            try:
                                if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                                    para_shading = paragraph._element.pPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                    for shd in para_shading:
                                        paragraph._element.pPr.remove(shd)
                                        highlighting_removed += 1
                            except:
                                pass
                            
                            for run in paragraph.runs:
                                # Apply all highlighting removal methods to footer text
                                try:
                                    if run.font.highlight_color is not None:
                                        run.font.highlight_color = None
                                        highlighting_removed += 1
                                except:
                                    pass
                                
                                try:
                                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                                        # Remove shading elements
                                        shading_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                        for elem in shading_elements:
                                            run._element.rPr.remove(elem)
                                            highlighting_removed += 1
                                        
                                        # Remove highlight elements
                                        highlight_elements = run._element.rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                                        for elem in highlight_elements:
                                            run._element.rPr.remove(elem)
                                            highlighting_removed += 1
                                except:
                                    pass
                except:
                    pass
        
        # Save the cleaned document
        doc.save(output_path)
        
        if highlighting_removed > 0:
            print(f"🎨 Removed highlighting from {highlighting_removed} text runs in DOCX")
        
        return True, f"Cleaned {highlighting_removed} highlighted sections"
        
    except ImportError:
        return False, "python-docx not available for highlighting removal"
    except Exception as e:
        return False, f"Error cleaning DOCX highlighting: {e}"

def filter_base64_from_csv(csv_content):
    """Filter out base64 logo data from CSV content to save API tokens."""
    lines = csv_content.split('\n')
    filtered_lines = []
    
    for line in lines:
        # Check if this line contains base64 logo data
        if ';Logo Base64 Data;_logo_base64_data;' in line:
            # Find where the actual base64 data starts (after "data:image/")
            if 'data:image/' in line and 'base64,' in line:
                # Split at base64, and keep everything before it + placeholder
                base64_start = line.find('base64,') + 7  # +7 for "base64,"
                if base64_start > 6:  # Valid base64 start found
                    before_base64 = line[:base64_start]
                    base64_data = line[base64_start:]
                    filtered_line = before_base64 + '[BASE64_DATA_REMOVED_FOR_API_EFFICIENCY]'
                    filtered_lines.append(filtered_line)
                    print(f"🖼️  FILTERED: Removed {len(base64_data):,} chars of base64 logo data to save API tokens!")
                    print(f"💰 API Cost Savings: ~${len(base64_data) * 0.000003:.2f} per request")
                else:
                    filtered_lines.append(line)
            else:
                # Fallback: try the old method with semicolon splitting
                parts = line.split(';', 4)
                if len(parts) >= 5:
                    # Keep the structure but replace data with placeholder
                    filtered_line = ';'.join(parts[:4]) + ';[BASE64_LOGO_DATA_REMOVED_FOR_API_EFFICIENCY]'
                    filtered_lines.append(filtered_line)
                    print(f"🖼️  Filtered out base64 logo data ({len(parts[4])} chars) to save API tokens")
                else:
                    filtered_lines.append(line)
        else:
            filtered_lines.append(line)
    
    return '\n'.join(filtered_lines)

def load_file_content(file_path):
    """Load content from various file types."""
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Handle DOCX files first
    if file_path.suffix.lower() == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                # Extract text while filtering out highlighted content
                clean_text = ""
                for run in paragraph.runs:
                    # Check if the run has highlighting
                    is_highlighted = False
                    try:
                        # Check for highlight color (python-docx way)
                        if run.font.highlight_color is not None:
                            is_highlighted = True
                    except:
                        pass
                    
                    # Also check for background color highlighting
                    try:
                        if hasattr(run.font, 'fill') and run.font.fill.fore_color.rgb is not None:
                            is_highlighted = True
                    except:
                        pass
                    
                    # Only add text if it's not highlighted
                    if not is_highlighted:
                        clean_text += run.text
                
                # Only add non-empty paragraphs
                if clean_text.strip():
                    content.append(clean_text)
            
            filtered_content = '\n'.join(content)
            print(f"📄 DOCX loaded: {len(filtered_content)} characters (highlighted text removed)")
            return filtered_content
            
        except ImportError:
            return f"[DOCX FILE: {file_path.name} - Install python-docx to read content]"
        except Exception as e:
            raise Exception(f"Error reading DOCX file {file_path}: {e}")
    
    # Handle JSON files (questionnaire responses)
    elif file_path.suffix.lower() == '.json':
        # Read JSON questionnaire answers and convert to CSV-like format for AI processing
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Convert JSON answers to CSV-like format for AI processing
        csv_lines = ['Question Number;Question Text;field;Response Type;User Response']
        
        for field, answer_data in json_data.items():
            if isinstance(answer_data, dict):
                question_number = answer_data.get('questionNumber', 0)
                question_text = answer_data.get('questionText', field)
                response_type = answer_data.get('responseType', 'text')
                value = answer_data.get('value', '')
                
                # Handle different value types
                if isinstance(value, dict) and 'data' in value:
                    # File upload - use filename or placeholder
                    value = value.get('name', 'uploaded_file')
                elif isinstance(value, (list, dict)):
                    value = str(value)
                
                csv_line = f"{question_number};{question_text};{field};{response_type};{value}"
                csv_lines.append(csv_line)
        
        content = '\n'.join(csv_lines)
        print(f"📊 Converted {len(json_data)} JSON answers to CSV format for AI processing")
        
        # No need to filter base64 data from JSON (it's already structured)
        return content
    
    # Handle CSV files with base64 filtering
    elif file_path.suffix.lower() == '.csv':
        # Read CSV as formatted text, but filter out base64 image data for API efficiency
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return filter_base64_from_csv(content)
    
    # Handle Markdown files
    elif file_path.suffix.lower() == '.md':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    # Default text reading for other file types
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

def extract_json_from_response(response_text):
    """Extract JSON content from Claude's response."""
    
    # Look for JSON blocks in the response
    json_patterns = [
        r'```json\n(.*?)\n```',
        r'```\n(\{.*?\})\n```',
        r'(\{[\s\S]*?"instructions"[\s\S]*?\})'
    ]
    
    for pattern in json_patterns:
        match = re.search(pattern, response_text, re.DOTALL | re.IGNORECASE)
        if match:
            content = match.group(1).strip()
            try:
                # Validate it's proper JSON
                parsed = json.loads(content)
                if 'metadata' in parsed and 'instructions' in parsed:
                    return content
            except json.JSONDecodeError:
                continue
    
    # Fallback: look for JSON structure starting with {
    lines = response_text.split('\n')
    json_start = -1
    brace_count = 0
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith('{') and json_start == -1:
            json_start = i
            brace_count += stripped.count('{') - stripped.count('}')
        elif json_start >= 0:
            brace_count += stripped.count('{') - stripped.count('}')
            if brace_count == 0:
                # Found complete JSON
                content = '\n'.join(lines[json_start:i+1])
                try:
                    parsed = json.loads(content)
                    if 'metadata' in parsed and 'instructions' in parsed:
                        return content
                except json.JSONDecodeError:
                    pass
                json_start = -1
                brace_count = 0
    
    raise ValueError("Could not extract valid JSON from Claude's response")

def validate_json_content(content):
    """Validate the generated JSON has the correct format."""
    
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON format: {e}")
    
    # Check required structure
    if 'metadata' not in parsed:
        raise ValueError("JSON must have 'metadata' section")
    
    if 'instructions' not in parsed:
        raise ValueError("JSON must have 'instructions' section")
    
    metadata = parsed['metadata']
    required_metadata = ['generated_timestamp', 'company_name', 'format_version', 'total_operations', 'generator']
    for field in required_metadata:
        if field not in metadata:
            raise ValueError(f"Missing required metadata field: {field}")
    
    # Check instructions structure
    instructions = parsed['instructions']
    if 'operations' not in instructions:
        raise ValueError("Instructions must have 'operations' array")
    
    operations = instructions['operations']
    if not isinstance(operations, list) or len(operations) == 0:
        raise ValueError("Operations must be a non-empty array")
    
    # Validate each operation
    for i, operation in enumerate(operations):
        # Check required fields
        required_fields = ['target_text', 'action', 'comment', 'comment_author']
        for field in required_fields:
            if field not in operation:
                raise ValueError(f"Operation {i+1} missing required field: {field}")
        
        # Check action is valid
        if operation['action'] not in ['replace', 'delete', 'comment', 'replace_with_logo']:
            raise ValueError(f"Operation {i+1} has invalid action: {operation['action']}")
        
        # Replacement field is optional for 'comment' and 'replace_with_logo' actions, required for others
        if operation['action'] in ['replace', 'delete']:
            if 'replacement' not in operation:
                raise ValueError(f"Operation {i+1} with action '{operation['action']}' missing required field: replacement")
        
        # Ensure replacement field exists (set to empty string if missing)
        if 'replacement' not in operation:
            operation['replacement'] = ''
    
    return True

def call_claude_api(prompt_content, questionnaire_content, policy_instructions_content, policy_content, api_key):
    """Call Claude Sonnet 4 API to generate JSON instructions."""
    
    # Import anthropic here when actually needed
    global anthropic
    if anthropic is None:
        try:
            import anthropic
        except ImportError:
            raise ImportError("anthropic package is required for API calls. Install it with: pip install anthropic")
    
    client = anthropic.Anthropic(api_key=api_key)
    
    # Construct the full prompt with the new JSON workflow
    full_prompt = f"""
{prompt_content}

---

## PROCESSING INSTRUCTIONS (Policy Document Specific Rules):
{policy_instructions_content}

---

## INPUT DATA FOR PROCESSING

### QUESTIONNAIRE RESPONSES (CSV FORMAT):
```csv
{questionnaire_content}
```

### POLICY DOCUMENT CONTENT (FOR REFERENCE):
```
{policy_content[:2000]}...
[Document truncated for API efficiency - full content will be processed by automation system]
```

---

Please analyze the questionnaire data and generate the complete JSON file for automated policy customization according to the processing instructions above.

CRITICAL: Your response must include a properly formatted JSON structure that follows the exact format specified in the processing instructions.
"""

    try:
        message = client.messages.create(
            model="claude-3-5-sonnet-20241022",  # Claude Sonnet model
            max_tokens=4000,
            temperature=0.1,  # Low temperature for consistent, accurate output
            messages=[{
                "role": "user",
                "content": full_prompt
            }]
        )
        
        return message.content[0].text
    
    except Exception as e:
        raise Exception(f"Claude API call failed: {e}")

def main():
    parser = argparse.ArgumentParser(description='AI Policy Processor - Generate JSON instructions with Claude Sonnet 4')
    parser.add_argument('--policy', required=True, help='Path to policy DOCX file')
    parser.add_argument('--questionnaire', help='Path to questionnaire CSV/JSON file')
    parser.add_argument('--questionnaire-env-data', action='store_true', help='Read questionnaire data from QUESTIONNAIRE_ANSWERS_DATA environment variable')
    parser.add_argument('--prompt', required=True, help='Path to AI prompt markdown file (prompt.md)')
    parser.add_argument('--policy-instructions', required=True, help='Path to policy processing instructions (updated_policy_instructions_v4.0.md)')
    parser.add_argument('--output', required=True, help='Output path for generated JSON file')
    parser.add_argument('--api-key', help='Claude API key (or set CLAUDE_API_KEY env var)')
    parser.add_argument('--skip-api', action='store_true', help='Skip API call and use existing JSON file (for testing/development)')
    
    args = parser.parse_args()
    
    # Validate questionnaire input
    if not args.questionnaire and not args.questionnaire_env_data:
        print("❌ Error: Either --questionnaire (file path) or --questionnaire-env-data must be provided!")
        sys.exit(1)
    
    if args.questionnaire and args.questionnaire_env_data:
        print("❌ Error: Cannot use both --questionnaire and --questionnaire-env-data at the same time!")
        sys.exit(1)
    
    # Check for skip API configuration
    skip_api_env = os.environ.get('SKIP_API_CALL', '').lower()
    skip_api = args.skip_api or skip_api_env in ['true', '1', 'yes', 'on']
    
    if skip_api:
        # Check if output file already exists
        if not Path(args.output).exists():
            print("❌ Error: --skip-api specified but output JSON file doesn't exist!")
            print(f"   Expected file: {args.output}")
            print("   Either run without --skip-api first, or provide an existing JSON file")
            sys.exit(1)
        
        print("🔄 SKIP_API_CALL enabled - Using existing JSON file for testing/development")
        print(f"📁 Using existing file: {args.output}")
        
        # Validate the existing JSON file
        try:
            with open(args.output, 'r', encoding='utf-8') as f:
                content = f.read()
            validate_json_content(content)
            
            # Show stats from existing file
            parsed = json.loads(content)
            operations_count = len(parsed['instructions']['operations'])
            company_name = parsed['metadata']['company_name']
            
            print(f"✅ Using existing JSON: {operations_count} operations for {company_name}")
            
            # Show operation types summary
            actions = {}
            for op in parsed['instructions']['operations']:
                action = op['action']
                actions[action] = actions.get(action, 0) + 1
            
            print("\n📋 Operations Summary (from existing file):")
            for action, count in actions.items():
                print(f"   {action}: {count} operations")
            
            print(f"\n💰 API call skipped - cost savings for testing/development!")
            return
            
        except Exception as e:
            print(f"❌ Error validating existing JSON file: {e}")
            sys.exit(1)
    
    # Get API key (only if not skipping API)
    api_key = args.api_key or os.environ.get('CLAUDE_API_KEY')
    if not api_key:
        print("❌ Error: Claude API key required!")
        print("   Set CLAUDE_API_KEY environment variable or use --api-key")
        print("   Get your key from: https://console.anthropic.com/")
        print("   Or use --skip-api to use existing JSON file for testing")
        sys.exit(1)
    
    print("🤖 AI Policy Processor Starting (JSON Mode)...")
    print(f"📋 Policy: {args.policy}")
    
    if args.questionnaire_env_data:
        print(f"📊 Questionnaire: Environment variable data")
    else:
        print(f"📊 Questionnaire: {args.questionnaire}")
        
    print(f"📝 Main Prompt: {args.prompt}")
    print(f"📜 Policy Instructions: {args.policy_instructions}")
    print(f"💾 Output: {args.output}")
    
    try:
        # Load input files
        print("\n📂 Loading input files...")
        policy_content = load_file_content(args.policy)
        
        # Load questionnaire content from environment variable or file
        if args.questionnaire_env_data:
            print("📊 Loading questionnaire data from environment variable...")
            env_data = os.environ.get('QUESTIONNAIRE_ANSWERS_DATA')
            if not env_data:
                print("❌ Error: QUESTIONNAIRE_ANSWERS_DATA environment variable not set!")
                sys.exit(1)
            
            try:
                # Parse and convert JSON to CSV-like format
                json_data = json.loads(env_data)
                csv_lines = ['Question Number;Question Text;field;Response Type;User Response']
                
                for field, answer_data in json_data.items():
                    if isinstance(answer_data, dict):
                        question_number = answer_data.get('questionNumber', 0)
                        question_text = answer_data.get('questionText', field)  # Use field as fallback
                        response_type = answer_data.get('responseType', 'text')
                        value = answer_data.get('value', '')
                        
                        # Handle different value types
                        if isinstance(value, dict) and 'data' in value:
                            # File upload - use filename or placeholder
                            value = value.get('name', 'uploaded_file')
                        elif field == '_logo_base64_data' and isinstance(value, str):
                            # Special case: logo base64 data - set correct metadata for existing filter
                            response_type = 'file_upload'  # Override for logo data
                            question_text = 'Logo Base64 Data'  # Override for logo data
                            # Note: base64 filtering will be applied by existing logic below
                        elif isinstance(value, (list, dict)):
                            value = str(value)
                        
                        csv_line = f"{question_number};{question_text};{field};{response_type};{value}"
                        csv_lines.append(csv_line)
                
                # Create CSV content and apply base64 filtering ONLY for API (keep original in env)
                raw_csv_content = '\n'.join(csv_lines)
                questionnaire_content = filter_base64_from_csv(raw_csv_content)
                print(f"📊 Converted {len(json_data)} JSON answers from environment to CSV format")
                print(f"🖼️  Note: Original base64 logo data preserved in environment for automation scripts")
                
            except json.JSONDecodeError as e:
                print(f"❌ Error: Invalid JSON in environment variable: {e}")
                sys.exit(1)
        else:
            print("📊 Loading questionnaire data from file...")
            questionnaire_content = load_file_content(args.questionnaire)
        
        prompt_content = load_file_content(args.prompt)
        policy_instructions_content = load_file_content(args.policy_instructions)
        
        print(f"✅ Policy loaded: {len(policy_content)} characters")
        print(f"✅ Questionnaire loaded: {questionnaire_content.count('Question')} questions detected")
        print(f"✅ Main prompt loaded: {len(prompt_content)} characters")
        print(f"✅ Policy instructions loaded: {len(policy_instructions_content)} characters")
        
        # Call Claude API
        print("\n🧠 Calling Claude Sonnet 4 API...")
        response = call_claude_api(prompt_content, questionnaire_content, policy_instructions_content, policy_content, api_key)
        
        print("✅ AI response received")
        
        # Extract JSON from response
        print("\n🔍 Extracting JSON from AI response...")
        content = extract_json_from_response(response)
        
        # Validate JSON
        print("✅ Validating JSON format...")
        validate_json_content(content)
        
        # Save output
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"\n🎉 SUCCESS! Generated JSON instructions:")
        print(f"📁 Saved to: {output_path}")
        
        # Show JSON stats
        parsed = json.loads(content)
        operations_count = len(parsed['instructions']['operations'])
        company_name = parsed['metadata']['company_name']
        
        print(f"📊 JSON Stats: {operations_count} operations for {company_name}")
        
        # Show operation types summary
        actions = {}
        for op in parsed['instructions']['operations']:
            action = op['action']
            actions[action] = actions.get(action, 0) + 1
        
        print("\n📋 Operations Summary:")
        for action, count in actions.items():
            print(f"   {action}: {count} operations")
        
        print(f"\n🚀 Next Step: Use this JSON with the tracked changes system!")
        print(f"   JSON → CSV Converter → GitHub Actions → Tracked Changes DOCX")
        
    except FileNotFoundError as e:
        print(f"❌ File Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
